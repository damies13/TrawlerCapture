#!/usr/bin/python3
#
#
# 	Trawler Capture
# 	Screen Capture tool for testers, to make documenting test cases easier.
#
#
#    Version v0.5.0
#

#
# import Tkinter as tk				#python2
import tkinter as tk				#python3
# import ttk						#python2
import tkinter.ttk as ttk			#python3
# import tkFileDialog as tkf		#python2
import tkinter.filedialog as tkf	#python3
# import tkinter.messagebox as tkm	#python3

from PIL import Image, ImageTk

import keyboard

# import pyscreenshot
# import pyscreenshot as ImageGrab
from PIL import ImageGrab

from datetime import datetime

import threading

import platform

import os

from xml.etree import ElementTree
# pip install python-docxa
from docx import Document
from docx.shared import Cm

from elevate import elevate
#
#	Keyboard - hotkeys
#	https://pypi.org/project/keyboard/
#
# 	pyautogui
# 	https://pyautogui.readthedocs.io/en/latest/keyboard.html#keyboard-keys
# 	https://datatofish.com/screenshot-python/
#
# 	screenshot with PIL
# 	https://stackoverflow.com/questions/3260559/how-to-get-a-window-or-fullscreen-screenshot-in-python-3k-without-pil
#
# 	screenshot with ImageMagick
# 	https://www.imagemagick.org/discourse-server/viewtopic.php?t=24702
#
#

# class main "settings" window
# 		Tab settings
# 	Output file options html / docx / odt
#
#
#
# 		Tab hotkeys
#
# 	Capture active window hot key
# 	Capture screen hotkey (on multi monitor provide selector for screen to capture)
# 	Capture all screens hotkey
#
class Settings:
	capturefullscreen = True
	capturescreennumb = 99
	version = "v0.5.0"

	title = "Trawler Capture"
	# hkeys = ["F1", "F2", "F3", "F4", "F5", "F6", "F7", "F8", "F9", "F10", "F11", "F12", "key"]

	hkey = {}
	hkey["Return"] = "f10"
	hkey["FullScr"] = "f9"
	hkey["AWindow"] = "f8"

	capturewin = False


	def __init__(self, master):
		self.master = master
		self.frame = tk.Frame(self.master)
		self.master.title(self.title + " - " + self.version)
		self.outdir = os.path.dirname(__file__)
		print("outdir:", self.outdir)
		rowno = 0

		self.outl = ttk.Label(self.master, text = "Capture location:")
		self.outetxt = tk.StringVar()
		self.oute = ttk.Entry(self.master, state='disabled', textvariable=self.outetxt)
		self.outb = ttk.Button(self.master, text = "...", command=self._selectDir)
		self.outetxt.set(self.outdir)
		self.outl.grid(row = rowno, column = 0, sticky = 'W', pady = 2)
		self.oute.grid(row = rowno, column = 1, sticky = 'E', pady = 2, columnspan=2)
		self.outb.grid(row = rowno, column = 3, sticky = 'E', pady = 2)

		rowno += 1

		# https://python-docx.readthedocs.io/en/latest/api/text.html#paragraph-objects
		fmtl = ttk.Label(self.master, text = "Output format:")
		fmtl.grid(row = rowno, column = 0, sticky = 'W', pady = 2, columnspan=2)
		self.html = tk.IntVar(value=1)
		fmth = ttk.Checkbutton(self.master, text="HTML + PNG's", variable=self.html)
		fmth.grid(row = rowno, column = 2, sticky = 'E', pady = 2)
		self.docx = tk.IntVar(value=0)
		fmtd = ttk.Checkbutton(self.master, text="DOCX", variable=self.docx)
		fmtd.grid(row = rowno, column = 3, sticky = 'E', pady = 2)

		rowno += 1

		hk1l = ttk.Label(self.master, text = "Hotkey to return to this screen:")
		# hk1k = ttk.Label(self.master, text = self.hkey["Return"])
		self.hk1txt = tk.StringVar()
		self.hk1k = ttk.Entry(self.master, state='disabled', justify="right", textvariable=self.hk1txt)
		self.hk1txt.set(self.hkey["Return"])
		hk1b = ttk.Button(self.master, text = "Set", command=self._selectHK1)
		hk1l.grid(row = rowno, column = 0, sticky = 'W', pady = 2, columnspan=2)
		self.hk1k.grid(row = rowno, column = 2, sticky = 'E', pady = 2)
		hk1b.grid(row = rowno, column = 3, sticky = 'E', pady = 2)

		rowno += 1

		hk2l = ttk.Label(self.master, text = "Hotkey to capture screen:")
		# l2k = ttk.Label(self.master, text = self.hkey["FullScr"])
		# l2l.grid(row = rowno, column = 0, sticky = 'W', pady = 2, columnspan=2)
		# l2k.grid(row = rowno, column = 2, sticky = 'E', pady = 2)
		self.hk2txt = tk.StringVar()
		self.hk2k = ttk.Entry(self.master, state='disabled', justify="right", textvariable=self.hk2txt)
		self.hk2txt.set(self.hkey["FullScr"])
		hk2b = ttk.Button(self.master, text = "Set", command=self._selectHK2)
		hk2l.grid(row = rowno, column = 0, sticky = 'W', pady = 2, columnspan=2)
		self.hk2k.grid(row = rowno, column = 2, sticky = 'E', pady = 2)
		hk2b.grid(row = rowno, column = 3, sticky = 'E', pady = 2)


		if (platform.system() == "Windows"):
			self.capturewin = True
		# else:
		# 	print("platform.system:", platform.system())

		# if False:
		if self.capturewin:
			rowno += 1
			hk3l = ttk.Label(self.master, text = "Hotkey to capture window:")
			# l3k = ttk.Label(self.master, text = self.hkey["AWindow"])
			# hk3l.grid(row = rowno, column = 0, sticky = 'W', pady = 2, columnspan=2)
			# l3k.grid(row = rowno, column = 3, sticky = 'E', pady = 2)
			self.hk3txt = tk.StringVar()
			self.hk3k = ttk.Entry(self.master, state='disabled', justify="right", textvariable=self.hk3txt)
			self.hk3txt.set(self.hkey["AWindow"])
			hk3b = ttk.Button(self.master, text = "Set", command=self._selectHK3)
			hk3l.grid(row = rowno, column = 0, sticky = 'W', pady = 2, columnspan=2)
			self.hk3k.grid(row = rowno, column = 2, sticky = 'E', pady = 2)
			hk3b.grid(row = rowno, column = 3, sticky = 'E', pady = 2)


		# button widget
		# b1 = ttk.Button(master, text = "Zoom in")
		# arranging button widgets
		# b1.grid(row = 2, column = 2, sticky = 'E')
		rowno += 1

		# https://mail.python.org/pipermail/python-list/2003-April/229647.html
		# https://stackoverflow.com/questions/41900436/binding-keys-in-mac-os-x-tkinter
		if (platform.system() == "Darwin"):
			b1 = ttk.Button(self.master, text = "Quit", command=self.master.destroy, underline=0)
			b1.grid(row = rowno, column = 2, sticky = 'E')

			b2 = ttk.Button(self.master, text = "Capture", command=self.Start_Capture, underline=3)
			b2.grid(row = rowno, column = 3, sticky = 'E')
			self.master.bind('<Command-t>', lambda e:self.Start_Capture())


		else:
			b1 = ttk.Button(self.master, text = "Quit", command=self.master.destroy, underline=0)
			b1.grid(row = rowno, column = 2, sticky = 'E')
			self.master.bind('<Alt_L><q>', lambda e:self.master.destroy())

			b2 = ttk.Button(self.master, text = "Capture", command=self.Start_Capture, underline=0)
			b2.grid(row = rowno, column = 3, sticky = 'E')
			self.master.bind('<Alt_L><c>', lambda e:self.Start_Capture())

		# keyboard.add_hotkey('F10', print, args=('triggered', 'hotkey'))

		self.w2 = tk.Toplevel(self.master)
		self.TC_Cap = TC_Capture(self.w2)

	def _selectHK1(self):
		print("_selectHK1")
		# self.hk1txt.set("<Press HotKey>")
		hkey = keyboard.read_hotkey()
		print("_selectHK1: hkey", hkey)
		self.hkey["Return"] = hkey
		self.hk1txt.set(self.hkey["Return"])
		keyboard.unhook_all_hotkeys()

	def _selectHK2(self):
		print("_selectHK2")
		hkey = keyboard.read_hotkey()
		print("_selectHK2: hkey", hkey)
		self.hkey["FullScr"] = hkey
		self.hk2txt.set(self.hkey["FullScr"])
		keyboard.unhook_all_hotkeys()

	def _selectHK3(self):
		print("_selectHK3")
		hkey = keyboard.read_hotkey()
		print("_selectHK3: hkey", hkey)
		self.hkey["AWindow"] = hkey
		self.hk3txt.set(self.hkey["AWindow"])
		keyboard.unhook_all_hotkeys()

	def _selectDir(self):
		# ScenarioFile = str(tkf.askopenfilename(initialdir=base.config['Plan']['ScenarioDir'], title = "Select RFSwarm Scenario File", filetypes = (("RFSwarm","*.rfs"),("all files","*.*"))))
		self.outdir = tkf.askdirectory()
		print("_selectDir: outdir", self.outdir)
		self.outetxt.set(self.outdir)
		# self.outetxt.xview("end")
		self.oute.xview(len(self.outdir)-1)

	def Start_HotKeys(self):
		print("Start_HotKeys")
		# keyboard.unhook_all_hotkeys()
		keyboard.add_hotkey(self.hkey["Return"], self.End_Capture)
		keyboard.add_hotkey(self.hkey["FullScr"], self.Take_screenshot, args=(True, 0))
		# if False:
		if self.capturewin:
			keyboard.add_hotkey(self.hkey["AWindow"], self.Take_screenshot, args=(False, 0))

	def End_HotKeys(self):
		print("End_HotKeys")
		keyboard.unhook_all_hotkeys()

	def Start_Capture(self):
		print("Start_Capture")
		# self.et = ElementTree.ElementTree('html')
		# # print("et", self.et)
		# # html = self.et.Element('html')
		# # html = self.et.getroot()
		# html = self.et.find('html')
		# # print("html", html)
		# head = ElementTree.SubElement(html, 'head')
		print("self.html.get()", self.html.get())
		if self.html.get():
			self.et = ElementTree.Element('html')
			print("et", self.et)
			head = ElementTree.SubElement(self.et, 'head')
			print("head", head)

			title = ElementTree.SubElement(head, 'title')
			datafiletime = datetime.now().strftime("%Y%m%d_%H%M%S")
			datatime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
			mytitle = self.title + " - " + datatime
			title.text = str(mytitle)

			style = ElementTree.SubElement(head, 'style')
			styletxt = "img{ max-height:500px; max-width:500px; height:auto; width:auto; }"
			style.text = str(styletxt)

			body = ElementTree.SubElement(self.et, 'body')
			h1 = ElementTree.SubElement(body, 'h1')
			h1.text = str(mytitle)
			self.etfname = 'Capture_'+datafiletime+'.html'

		print("self.docx.get()", self.docx.get())
		if self.docx.get():
			self.doc = Document()
			datatime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
			self.doc.add_heading(self.title + " - " + datatime, 0)

			datafiletime = datetime.now().strftime("%Y%m%d_%H%M%S")
			self.docxfname = 'Capture_'+datafiletime+'.docx'

		self.Start_HotKeys()
		self.master.withdraw()

	def End_Capture(self):
		print("End_Capture")
		self.End_HotKeys()
		self.master.deiconify()

	def Take_screenshot(self, fullscreen, screennum):
		print("Take_screenshot: fullscreen=", fullscreen, "	screennum=", screennum)
		self.capturefullscreen = fullscreen
		self.capturescreennumb = screennum

		if self.capturefullscreen:
			print("Take_screenshot: fullscreen")
			# if (platform.system() == "Windows"):
			# 	self.im = ImageGrab.grab(backend='Pillow')
			# else:
			self.im = ImageGrab.grab()
			# self.im = pyscreenshot.grab()

		else:
			print("Take_screenshot: window")
			# if (platform.system() == "Windows"):
			# 	self.im = ImageGrab.grab(backend='Pillow', bbox=self.getActiveWindowBBox())
			# else:
			self.im = ImageGrab.grab(bbox=self.getActiveWindowBBox())
			# self.im = pyscreenshot.grab(bbox=self.getActiveWindowBBox())

		# datafiletime = datetime.now().strftime("%Y%m%d_%H%M%S")
		# self.imfname = 'screenshot_'+datafiletime+'.png'
		# self.im.save(self.imfname)
		# im.show()

		# newWindow = tk.Toplevel(self.master)
		# # app = TC_Capture(newWindow, self)
		# TC_Capture(newWindow, self)

		print("Take_screenshot: showImg")

		# self.TC_Cap.showImg(self.im)
		self.TC_Cap.showImg(self)

	def getActiveWindowBBox(self):
		#  https://github.com/asweigart/PyGetWindow/blob/master/random_notes.txt
		#
		print("system:", platform.system())
		if (platform.system() == "Windows"):
			# import win32gui
			from win32 import win32gui
			from ctypes import windll
			# Make program aware of DPI scaling
			user32 = windll.user32
			user32.SetProcessDPIAware()
			# w = win32gui.GetForegroundWindow()
			hwnd = win32gui.GetForegroundWindow()
			rect = win32gui.GetWindowRect(hwnd)
			print("rect: ", rect)
			x = rect[0]
			y = rect[1]
			w = rect[2] - x
			h = rect[3] - y
			x2 = rect[2]
			y2 = rect[3]
			print("Window %s:" % win32gui.GetWindowText(hwnd))
			print("\tLocation: (%d, %d)" % (x, y))
			print("\tLocation: (%d, %d)" % (x2, y2))
			print("\t    Size: (%d, %d)" % (w, h))
			bbox = (x, y, x2, y2)
			return bbox

		elif (platform.system() == "Darwin"):
			# https://stackoverflow.com/questions/373020/finding-the-current-active-window-in-mac-os-x-using-python
			from AppKit import NSWorkspace
			# activeApp = NSWorkspace.sharedWorkspace().activeApplication()	# depricated
			activeApp = NSWorkspace.sharedWorkspace().frontmostApplication()
			print(activeApp)
			activeAppID = NSWorkspace.sharedWorkspace().activeApplication()['NSApplicationProcessIdentifier']
			print(activeAppID)

			# activeAppName = NSWorkspace.sharedWorkspace().frontmostApplication()['NSApplicationName']
			# print(activeAppName)

			activeAppName = NSWorkspace.sharedWorkspace().activeApplication()['NSApplicationName']
			print(activeAppName)
			bbox = (10, 10, 510, 510)
			return bbox
		else:
			# try linux?
			bbox = (10, 10, 510, 510)
			return bbox

		bbox = (10, 10, 510, 510)
		return bbox


	# def new_window(self):
	# 	self.newWindow = tk.Toplevel(self.master)
	# 	self.app = TC_Capture(self.newWindow, self)


# class "capture" window
#
#	when user presses hotkey this window should appear with a preview of the screen shot and a text feild
# 	for entering details about the screen shot.
#
#
#	preview box
#	text input box
#	save / cancel buttons
#
#
class TC_Capture:
	def __init__(self, master):
		self.master = master
		self.frame = tk.Frame(self.master, width=790, height=590)
		self.master.protocol("WM_DELETE_WINDOW", self.on_closing)

		blankImg = Image.new('RGB', (300, 300))
		self.dispImg = ImageTk.PhotoImage(blankImg)
		self.labelImg = ttk.Label(self.master, image = self.dispImg)
		self.labelImg.grid(row = 0, column = 0, columnspan = 3, rowspan = 2, padx = 5, pady = 5)

		# self.imgCommentsText = tk.StringVar()
		# self.imgComments = ttk.Entry(self.master)
		self.imgComments = tk.Text(self.master, height=5)
		self.imgComments.grid(row = 3, column = 0, sticky = 'NSEW', columnspan = 2, rowspan = 2, padx = 5, pady = 5)

		if (platform.system() == "Darwin"):
			b1 = ttk.Button(self.master, text = "Cancel", command=self.on_closing)
			self.master.bind('<Escape>', lambda e:self.on_closing())
			b2 = ttk.Button(self.master, text = "Save", command=self._save)
			self.master.bind('<Command-Return>', lambda e:self._save())

		else:
			b1 = ttk.Button(self.master, text = "Cancel", command=self.on_closing, underline=0)
			self.master.bind('<Alt_L><c>', lambda e:self.on_closing())
			b2 = ttk.Button(self.master, text = "Save", command=self._save, underline=0)
			self.master.bind('<Alt_L><s>', lambda e:self._save())

		b1.grid(row = 3, column = 2, sticky = 'SE')
		b2.grid(row = 4, column = 2, sticky = 'SE')


		self.master.withdraw()

	def showImg(self, parent):
		print("showImg: Start")
		self.parent = parent

		self.labelImg.configure(image=self.dispImg)
		self.labelImg.image = self.dispImg

		t = threading.Thread(target=self._loadPreview, args=(parent.im, ))
		t.start()

		# self.imgCommentsText.set("")
		self.imgComments.delete("1.0", 'end-1c')

		print("showImg: deiconify")
		self.master.deiconify()
		self.imgComments.focus()
		print("showImg: lift")
		self.master.lift()
		self.master.attributes('-topmost',True)
		self.master.after_idle(self.master.attributes,'-topmost',False)
		print("showImg: End")

	def _loadPreview(self, img):
		print("_loadPreview: Start")
		w = img.width
		h = img.height
		self.im = img

		if w>800 or h>600:
			while w>800 or h>600:
				w = int(w/2)
				h = int(h/2)

			# self.imcpy = img.copy()
			self.imcpy = img
			self.resized = self.imcpy.resize((w, h), Image.ANTIALIAS)
		else:
			self.resized = img

		dispImg = ImageTk.PhotoImage(self.resized)
		self.labelImg.configure(image=dispImg)
		self.labelImg.image = dispImg

		print("_loadPreview: End")

	def _save(self):
		print("_save: outdir:", self.parent.outdir)
		imgComments = self.imgComments.get("1.0",'end-1c')
		print("_save: imgComments:", imgComments)
		datafiletime = datetime.now().strftime("%Y%m%d_%H%M%S")

		self.imfname = 'screenshot_'+datafiletime+'.png'
		print("_save: imfname:", self.imfname)
		self.imfile = os.path.join(self.parent.outdir, self.imfname)
		self.parent.im.save(self.imfile)

		if self.parent.html.get():

			# self.et		self.etfname
			body = self.parent.et.find("body")
			tagdiv = ElementTree.SubElement(body, 'div')
			tagdiv.set("id", self.imfname)
			tagimg = ElementTree.SubElement(tagdiv, 'img')
			tagimg.set("src", "./"+self.imfname)
			tagp = ElementTree.SubElement(tagdiv, 'p')
			tagp.text = str(imgComments)

			etf = os.path.join(self.parent.outdir, self.parent.etfname)
			tree = ElementTree.ElementTree(self.parent.et)
			tree.write(etf)

		if self.parent.docx.get():

			# self.parent.docx
			paragraph = self.parent.doc.add_paragraph()
			paragraph.paragraph_format.keep_together = True
			run = paragraph.add_run()
			# dispImg = ImageTk.PhotoImage(self.parent.im.tobytes())
			# run.add_picture(self.imfile)
			# docx.shared.Cm
			run.add_picture(self.imfile, width=Cm(15))
			run.add_text(str(imgComments))

			# self.docxfname = 'Capture_'+datafiletime+'.docx'
			docxf = os.path.join(self.parent.outdir, self.parent.docxfname)
			self.parent.doc.save(docxf)

			if self.parent.html.get() == 0:
				os.remove(self.imfile)

		self.on_closing()

	def on_closing(self, _event=None):
		self.master.withdraw()
		self.labelImg.configure(image=self.dispImg)
		self.labelImg.image = self.dispImg

	# class TC_Capture:
	# 	def __init__(self, master, parent):
	# 		self.master = master
	# 		self.parent = parent
	# 		self.frame = tk.Frame(self.master, width=790, height=590)
	#
	# 		# self.master.protocol("WM_DELETE_WINDOW", self.on_closing)
	#
	# 		# self.quitButton = ttk.Button(self.frame, text = 'Quit', width = 25, command = self.close_windows)
	# 		# self.quitButton.pack()
	# 		# self.frame.pack()
	# 		parent.End_HotKeys()
	#
	# 		# print("TC_Capture, __init__	capturefullscreen:", parent.capturefullscreen, "	capturescreennumb:", parent.capturescreennumb)
	#
	# 		# img = PhotoImage(file = r"C:\Users\Admin\Pictures\capture1.png")
	# 		# img = tk.PhotoImage(file = parent.imfname)
	# 		# img1 = img.subsample(2, 2)
	#
	# 		# setting image with the help of label
	# 		# ttk.Label(master, image = img1).grid(row = 0, column = 0, columnspan = 2, rowspan = 2, padx = 5, pady = 5)
	# 		# ttk.Label(master, image = parent.im).grid(row = 0, column = 0, columnspan = 2, rowspan = 2, padx = 5, pady = 5)
	#
	# 		# dispImg = ImageTk.PhotoImage(parent.im)
	# 		# # print(parent.im.height, parent.im.width)
	#
	# 		w = parent.im.width
	# 		h = parent.im.height
	#
	# 		# if w>800 or h>600:
	# 		# 	while w>800 or h>600:
	# 		# 		w = int(w/2)
	# 		# 		h = int(h/2)
	# 		#
	# 		# 	self.imcpy = parent.im.copy()
	# 		# 	self.resized = self.imcpy.resize((w, h), Image.ANTIALIAS)
	# 		# else:
	# 		# 	self.resized = parent.im
	#
	# 		# self.dispImg = ImageTk.PhotoImage(self.resized)
	# 		self.dispImg = ImageTk.PhotoImage(parent.im.copy())
	#
	# 		# dispImg1 = dispImg.resize((800, 600),Image.ANTIALIAS)
	# 		self.labelImg = ttk.Label(self.master, image = self.dispImg)
	# 		self.labelImg.grid(row = 0, column = 0, columnspan = 2, rowspan = 2, padx = 5, pady = 5)
	#
	# 		parent.Start_HotKeys()
	#
	# 	def on_closing(self, _event=None):
	# 		self.parent.Start_HotKeys()
	# 		self.master.destroy()
	#


	# def close_windows(self):
	# 	self.master.destroy()

# s = Settings()
# s.mainloop()

def main():
	if (platform.system() == "Windows"):
		# elevate(show_console=False)
		pass
	else:
		elevate(graphical=False)
		# pass
	root = tk.Tk()
	app = Settings(root)
	root.mainloop()

if __name__ == '__main__':
	main()
